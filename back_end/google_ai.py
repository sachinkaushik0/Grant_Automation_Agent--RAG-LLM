import time
import os
import google.generativeai as genai
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging

os.environ["GOOGLE_API_KEY"] = "AIzaSyBfXHGv_LCiLZSX9GMfjqulk10WTBV5DdA"

genai.configure(api_key=os.environ["GOOGLE_API_KEY"])


# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

generation_config = {
  "temperature": 1,
  "top_p": 0.95,
  "top_k": 40,
  "max_output_tokens": 8192,
  "response_mime_type": "text/plain",
}

model = genai.GenerativeModel(
  model_name="gemini-2.0-flash",
  generation_config=generation_config,
)

chat_session = model.start_chat(
  history=[
  ]
)

def generate_section_content(prompt):
    """Generate content using Falcon-7B."""
    generated_text = chat_session.send_message(prompt)
    return generated_text


def generate_proposal(grant_details):
    """Generate a proposal for the specific grant."""

    # Sample applicant details (can be replaced with user input)
    applicant_name = "John Doe"
    address = "123 Main St, Toronto, ON, Canada"
    email = "john.doe@example.com"
    phone_number = "+1 234 567 8901"

    # Proposal Template Sections
    sections = {
        "Applicant Name": applicant_name,
        "Address": address,
        "Email": email,
        "Phone Number": phone_number,
        "Grant Name": grant_details["program_name"],
        "Purpose of the Grant": generate_section_content(
            f"Describe the purpose of the grant based on grant description {grant_details['full_text']} grant:"),
        "Personal Background & Qualifications": generate_section_content(
            f"Describe the applicant's background and qualifications for the {grant_details['program_name']} grant:"),
        "Project Description": generate_section_content(
            f"Describe the project for the {grant_details['program_name']}  based on the below grant description {grant_details['full_text']} :"),
        "Expected Outcomes & Impact": generate_section_content(
            f"Describe the expected outcomes and impact of the {grant_details['program_name']} grant based on the description {grant_details['full_text']} :"),
        "Budget & Financial Requirements": generate_section_content(
            f"Provide a budget and financial breakdown for the {grant_details['program_name']} grant based on the description {grant_details['full_text']} :"),
        "Conclusion & Commitment": generate_section_content(
            f"Provide a conclusion and commitment for the {grant_details['program_name']} grant based on the description {grant_details['full_text']}  :")
    }
    return sections


def save_proposal_to_docx(sections, file_name):
    """Save the proposal to a DOCX file."""
    doc = Document()
    doc.add_heading("Grant Proposal", level=1)

    for section, content in sections.items():
        doc.add_heading(section, level=2)
        if hasattr(content, 'text'):
            content = content.text  # Extract text from the response object
        para = doc.add_paragraph(content)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.save(file_name)
    logger.info(f"Proposal saved as {file_name}")


def main():
    # Grant details (hardcoded for now)
    grant_details = {
        "program_name" : "Alliance Advantage",
        "program_source" : "NSERC",
        "program_type" : "Funding/Grant, Partnering and Collaboration",
        "program_target" : "Scientific Research, Job Training and Employment, Workforce, Development, Labor",
        "main_industry" : "Public Administration",
        "country": "Canada",
        "target_audience": "null",
        "min_funding": 20000,
        "max_funding": 1000000,
        "amount" : 20000,
        "unit": "CAD",
        "url": "https://www.nserc-crsng.gc.ca/Innovate-Innover/alliance-alliance/index_eng.asp",
        "full_text" : "Alliance Advantage\n\n\nWho?\nUniversity researchers collaborating with private-sector, public-sector or not-for-profit organizationsCollege faculty can participate as co-applicants\n\n\nHow much?\n$20,000 to $1 million per year\n\n\nHow long?\n1 to 5 years\n\n\n\n\n\n\n\n\n\nOverview\nAlliance grants encourage university researchers to collaborate with partner organizations, which can be from the private, public or not-for-profit sectors. These grants support research projects led by strong, complementary, collaborative teams that will generate new knowledge and accelerate the application of research results to create benefits for Canada.\nAlliance Advantage grants (formerly Alliance cost-sharing option 1) are for partner-driven projects. They fund projects focused on the partners’ goals, with at least one partner sharing in the costs of research. Research supported by Alliance Advantage grants will:\n\ngenerate new knowledge and/or technology to address complex challenges\ncreate economic, environmental and/or other societal benefits\ncontribute to Canada’s long-term competitiveness\nsupport public policy\ntrain new researchers in areas that are important to Canada and partner organizations\ndraw on diverse perspectives and skill sets to accelerate the translation and application of research results\n\nFor projects with societal impact as the main driver (formerly cost-sharing option 2), see Alliance Society.\n\nWho can apply?\nIf you are a Canadian university researcher who is eligible to receive NSERC funds, you can apply on your own or as a team with co-applicants who are also eligible academic researchers. You must have at least one partner organization (in the private, public or not-for-profit sector) that can be recognized for cost sharing (see Alliance Advantage: Partner organizations), but you may include other partner organizations who play an important role in your research project, whether or not their cash contributions are recognized for cost sharing.\nAlliance grants support projects of varying scale and complexity, from short-term smaller projects involving one researcher to long-term projects involving researchers across several universities, and from one-on-one collaborations with one partner organization directly involved in the research to projects involving many partner organizations across multiple sectors.\n\nEquity, diversity and inclusion\nNSERC is acting on the evidence that achieving a more equitable, diverse and inclusive Canadian research enterprise is essential to creating the excellent, innovative and impactful research necessary to advance knowledge and understanding, and to respond to local, national and global challenges. This principle informs the commitments described in the Tri-agency statement on equity, diversity and inclusion (EDI) and is aligned with the objectives of the Tri-agency EDI action plan.\nExcellent research considers EDI both in the research environment (forming a research team, student training) and in the research process. For Alliance grants, EDI considerations are currently evaluated in the training, mentorship and professional development opportunities for students and trainees. The aim is to remove barriers to the recruitment and promote full participation of individuals from underrepresented groups, including women, Indigenous Peoples (First Nations, Inuit, and Métis), persons with disabilities, members of visible minority/racialized groups and members of 2SLGBTQI+ communities. Applicants are encouraged to increase the inclusion and advancement of underrepresented groups as one way to enhance excellence in research and training. For additional guidance, applicants should refer to Alliance grants: Equity, diversity and inclusion in your training plan and the NSERC guide on integrating equity, diversity and inclusion considerations in research.\n\nNational Security Guidelines for Research Partnerships\nTo ensure the Canadian research ecosystem is as open as possible and as safeguarded as necessary, the Government of Canada has introduced the National Security Guidelines for Research Partnerships to integrate national security considerations into the development, evaluation and funding of research partnerships. These guidelines provide a framework through which researchers, research institutions and Canada’s granting agencies can undertake consistent, risk-targeted due diligence to identify and mitigate potential national security risks linked to research partnerships.\nThe National Security Guidelines for Research Partnerships apply to Alliance grant applications involving one or more partner organizations from the private sector, including when they participate alongside other partner organizations from the public and/or not-for-profit sectors. For such partnerships, you and your post-secondary institution are required to complete a risk assessment form for your research project and submit it as an integral part of your Alliance application.\nRefer to How to apply and the instructions for completing an Alliance grant application for further information.\n\nWhat's next?\r\n\t\t\t\t\t\t\tPartner organizations\n\n\n\n\n\nApply now\nSign up\nSign up to receive email updates with the latest information on NSERC Alliance grants.\nContactalliance@nserc-crsng.gc.ca\r\n1-855-275-2861\nApplication deadline\r\nNo deadline \n\n\nSign up\nSign up to receive the latest information on NSERC Alliance Grants\n\n\n\n\n\n\n\n\n\n\n\n\n\nContact Newsletter\nGet highlights of things happening at NSERC delivered to your email inbox. View all Newsletters\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\nTOP OF PAGE\n\n\n\n\n\n\n\nModified: 2024-04-17\n\n\n\n\n\n\n\n\n\nAlliance Advantage\n\n\nWho?\nUniversity researchers collaborating with private-sector, public-sector or not-for-profit organizationsCollege faculty can participate as co-applicants\n\n\nHow much?\n$20,000 to $1 million per year\n\n\nHow long?\n1 to 5 years\n\n\n\n\n\n\n\n\n\nOverview\nAlliance grants encourage university researchers to collaborate with partner organizations, which can be from the private, public or not-for-profit sectors. These grants support research projects led by strong, complementary, collaborative teams that will generate new knowledge and accelerate the application of research results to create benefits for Canada.\nAlliance Advantage grants (formerly Alliance cost-sharing option 1) are for partner-driven projects. They fund projects focused on the partners’ goals, with at least one partner sharing in the costs of research. Research supported by Alliance Advantage grants will:\n\ngenerate new knowledge and/or technology to address complex challenges\ncreate economic, environmental and/or other societal benefits\ncontribute to Canada’s long-term competitiveness\nsupport public policy\ntrain new researchers in areas that are important to Canada and partner organizations\ndraw on diverse perspectives and skill sets to accelerate the translation and application of research results\n\nFor projects with societal impact as the main driver (formerly cost-sharing option 2), see Alliance Society.\n\nWho can apply?\nIf you are a Canadian university researcher who is eligible to receive NSERC funds, you can apply on your own or as a team with co-applicants who are also eligible academic researchers. You must have at least one partner organization (in the private, public or not-for-profit sector) that can be recognized for cost sharing (see Alliance Advantage: Partner organizations), but you may include other partner organizations who play an important role in your research project, whether or not their cash contributions are recognized for cost sharing.\nAlliance grants support projects of varying scale and complexity, from short-term smaller projects involving one researcher to long-term projects involving researchers across several universities, and from one-on-one collaborations with one partner organization directly involved in the research to projects involving many partner organizations across multiple sectors.\n\nEquity, diversity and inclusion\nNSERC is acting on the evidence that achieving a more equitable, diverse and inclusive Canadian research enterprise is essential to creating the excellent, innovative and impactful research necessary to advance knowledge and understanding, and to respond to local, national and global challenges. This principle informs the commitments described in the Tri-agency statement on equity, diversity and inclusion (EDI) and is aligned with the objectives of the Tri-agency EDI action plan.\nExcellent research considers EDI both in the research environment (forming a research team, student training) and in the research process. For Alliance grants, EDI considerations are currently evaluated in the training, mentorship and professional development opportunities for students and trainees. The aim is to remove barriers to the recruitment and promote full participation of individuals from underrepresented groups, including women, Indigenous Peoples (First Nations, Inuit, and Métis), persons with disabilities, members of visible minority/racialized groups and members of 2SLGBTQI+ communities. Applicants are encouraged to increase the inclusion and advancement of underrepresented groups as one way to enhance excellence in research and training. For additional guidance, applicants should refer to Alliance grants: Equity, diversity and inclusion in your training plan and the NSERC guide on integrating equity, diversity and inclusion considerations in research.\n\nNational Security Guidelines for Research Partnerships\nTo ensure the Canadian research ecosystem is as open as possible and as safeguarded as necessary, the Government of Canada has introduced the National Security Guidelines for Research Partnerships to integrate national security considerations into the development, evaluation and funding of research partnerships. These guidelines provide a framework through which researchers, research institutions and Canada’s granting agencies can undertake consistent, risk-targeted due diligence to identify and mitigate potential national security risks linked to research partnerships.\nThe National Security Guidelines for Research Partnerships apply to Alliance grant applications involving one or more partner organizations from the private sector, including when they participate alongside other partner organizations from the public and/or not-for-profit sectors. For such partnerships, you and your post-secondary institution are required to complete a risk assessment form for your research project and submit it as an integral part of your Alliance application.\nRefer to How to apply and the instructions for completing an Alliance grant application for further information.\n\nWhat's next?\r\n\t\t\t\t\t\t\tPartner organizations\n\n\n\n\n\nApply now\nSign up\nSign up to receive email updates with the latest information on NSERC Alliance grants.\nContactalliance@nserc-crsng.gc.ca\r\n1-855-275-2861\nApplication deadline\r\nNo deadline \n\n\nSign up\nSign up to receive the latest information on NSERC Alliance Grants\n\n\n\n\n\n\n\n\n\n\n\n\nSign up\nSign up to receive the latest information on NSERC Alliance Grants\n\n\n\n\n\n\n\n\nContact Newsletter\nGet highlights of things happening at NSERC delivered to your email inbox. View all Newsletters\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\nTOP OF PAGE\n\n\n\n\n\n\n\nModified: 2024-04-17"
    }

    # Generate proposal sections
    sections = generate_proposal(grant_details)

    # Save proposal to file
    file_name = f"{grant_details['program_name'].replace(' ', '_')}_Proposal.docx"
    save_proposal_to_docx(sections, file_name)


if __name__ == "__main__":
    main()